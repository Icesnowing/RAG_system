"""
这是一个完整的本地RAG知识库系统
使用方式：
1.将文档放入./RAG_files文件夹下
2.运行rag_system.py文件，首次运行会自动建库
3.向AI提问，它会基于你的文档回答问题
"""

import os
import hashlib
import threading
import logging
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Set
# 导入文档加载器
from langchain_community.document_loaders import(
    PyPDFLoader,
    TextLoader,
    DirectoryLoader,
)

# 导入文本切分器
from langchain_text_splitters import RecursiveCharacterTextSplitter, MarkdownHeaderTextSplitter
#导入ollama模型的文本嵌入器、对话模式
from langchain_ollama import OllamaEmbeddings, ChatOllama
# 导入向量数据库
from langchain_chroma import Chroma
#导入对话提示词模板
from langchain_core.prompts import ChatPromptTemplate
#导入langchain参数透传
from langchain_core.runnables import RunnablePassthrough
from langchain_core.runnables import RunnableLambda
from langchain_unstructured import UnstructuredLoader
import re
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from langchain_core.documents import Document as LangchainDocument
from langchain_community.retrievers import BM25Retriever
from langchain_classic.retrievers import EnsembleRetriever
from langchain_classic.chains.retrieval import create_retrieval_chain
from langchain_classic.chains.combine_documents import create_stuff_documents_chain
# ================== 多路召回 + 重排依赖 ==================
from sentence_transformers import CrossEncoder
from rank_bm25 import BM25Okapi
import numpy as np
import jieba

from langchain_core.callbacks import StreamingStdOutCallbackHandler
from concurrent.futures import ThreadPoolExecutor, as_completed
from tqdm import tqdm
import pandas as pd
import argparse
import json
import time
from dataclasses import dataclass, field
from datetime import datetime
from functools import wraps

# ================== [OPT-09] 异常重试装饰器 ==================
def retry_on_failure(max_retries: int = 2, delay: float = 1.0, fallback_value=None):
    """
    对模型调用进行自动重试，防止偶发网络/模型错误导致流程中断。
    原理：LLM推理是概率性的，偶发超时或OOM不应中断批量评估；
          重试+退避可以恢复绝大多数瞬态错误。
    """
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            last_exception = None
            for attempt in range(max_retries + 1):
                try:
                    return func(*args, **kwargs)
                except Exception as e:
                    last_exception = e
                    if attempt < max_retries:
                        time.sleep(delay * (attempt + 1))
                    else:
                        if fallback_value is not None:
                            return fallback_value
                        raise last_exception
            return fallback_value
        return wrapper
    return decorator

#==========配置项============
DOCS_DIR = '../RAG_files'   #文档目录
CHROMA_DB_DIR = '../chroma_db'   #向量数据库目录
MANIFEST_PATH = os.path.join(CHROMA_DB_DIR, 'document_manifest.json')  #文档清单文件
EMBED_MODEL = 'nomic-embed-text-v2-moe'
CHAT_MODEL = 'qwen2.5:7b'
CHUNK_SIZE = 500
CHUNK_OVERLAP = 50
TOP_K = 10
RERANK_TOP_N = 5     # 重排后给LLM的数量（3条通常足够，5条会显著增加生成耗时）
CLEAN_DOCS_DIR = '../RAG_files_clean'
DEFAULT_TEST_SET = 'rag_testsets.csv'
EVAL_RESULTS_DIR = 'eval_results'

SIMILARITY_THRESHOLD = 0.3          # 检索相似度阈值：低于该值直接判定为无有效资料
ENABLE_KEYWORD_AUGMENT = False      # 是否启用关键词增强
USE_LLM_FOR_KEYWORDS = False       # 是否使用LLM提取关键词（False则使用规则提取）
KEYWORD_EXTRACTION_MODEL = CHAT_MODEL  # 提取关键词使用的模型
MAX_KEYWORDS = 8                    # 最多提取多少个关键词
# 多Query检索配置
ENABLE_MULTI_QUERY = False           # 是否启用多Query检索
MULTI_QUERY_NUM = 3                 # 生成的额外Query数量
MULTI_QUERY_MODEL = CHAT_MODEL      # 生成多Query使用的模型
# ================== 文档压缩配置 ==================
ENABLE_DOC_COMPRESSION = False       # 是否启用文档压缩（LLM抽取式压缩）
COMPRESSION_TARGET_LENGTH = 300     # 压缩目标长度（字符数），每个文档压缩后的目标长度
COMPRESSION_RATIO = 0.5             # 压缩比例（如果设置此值，会覆盖target_length）
COMPRESSION_MODEL = CHAT_MODEL      # 压缩使用的模型
COMPRESSION_TEMPERATURE = 0.1       # 压缩时的温度（低温度保证提取准确性）
MAX_WORKERS = 4 #并发处理数
#=====================

# [OPT-14] 耗时优化配置
RERANK_CANDIDATE_LIMIT = 15      # RRF融合后送入重排模型的最大候选数（减少CrossEncoder计算量）
RERANK_BATCH_SIZE = 16           # CrossEncoder批推理大小
ENABLE_PARALLEL_RECALL = True    # 三路召回是否并行执行
DEDUP_METHOD = "text"            # 去重方式: "text"(快速字符相似度) 或 "embedding"(慢但更准)
TEXT_DEDUP_THRESHOLD = 0.75      # 文本去重相似度阈值

# [OPT-15] 模型冷启动与推理加速（190s 主要来自 Ollama 首次加载 + LLM 生成）
ENABLE_RERANK = True             # False 时跳过重排，仅用 RRF Top-N（约省 5-15s/次）
PRELOAD_MODELS = True            # 启动时预热 embedding/rerank/chat，避免首次问答卡顿
OLLAMA_KEEP_ALIVE = "60m"        # 保持 Ollama 模型常驻内存，避免每次重新加载
CHAT_NUM_CTX = 4096              # 降低上下文窗口分配（原4096），加快推理
CHAT_NUM_PREDICT = 300           # 限制最大生成长度，避免冗长输出拖慢响应
CHAT_STREAMING = True           # 流式输出会略增开销；交互模式仍可见完整回答
ENABLE_PERF_LOG = False          # True 时打印各阶段耗时明细
FAST_MODE = False                # True = 跳过重排，检索更快但精度略降

# [OPT-08扩展] 上下文窗口管理
MAX_CONTEXT_TOKENS = 3000        # 送给LLM的最大context token数（越小生成越快）
DEDUP_THRESHOLD = 0.85           # chunk去重相似度阈值 [OPT-07]（仅embedding模式使用）
ENABLE_DEDUP = True              # 是否启用chunk去重 [OPT-07]

# 预热jieba分词，避免首次查询时的冷启动延迟
jieba.initialize()

@dataclass
class EvalResult:
    """评估结果数据类"""
    question: str
    ground_truth: str
    retrieved_answer: str
    retrieved_contexts: list[str]
    question_type: str
    difficulty: str
    
    # 评估指标
    answer_relevance: float = 0.0
    answer_accuracy: float = 0.0
    faithfulness: float = 0.0
    context_recall: float = 0.0

    # [OPT-04新增] 子维度指标，便于细粒度分析
    fact_coverage: float = 0.0        # 事实覆盖度（ground_truth中的关键事实被答案覆盖的比例）
    hallucination_rate: float = 0.0   # 幻觉率（答案中无法从context验证的比例）

    # 性能指标
    retrieval_time_ms: float = 0.0
    generation_time_ms: float = 0.0
    total_time_ms: float = 0.0
    
    # 检索质量
    retrieved_count: int = 0
    has_relevant_context: bool = False
    def to_dict(self):
        return {
            'question': self.question,
            'ground_truth': self.ground_truth,
            'retrieved_answer': self.retrieved_answer[:500] + '...' if len(self.retrieved_answer) > 500 else self.retrieved_answer,
            'question_type': self.question_type,
            'difficulty': self.difficulty,
            'answer_relevance': round(self.answer_relevance, 4),
            'answer_accuracy': round(self.answer_accuracy, 4),
            'faithfulness': round(self.faithfulness, 4),
            'context_recall': round(self.context_recall, 4),
            'fact_coverage': round(self.fact_coverage, 4),
            'hallucination_rate': round(self.hallucination_rate, 4),
            'retrieval_time_ms': round(self.retrieval_time_ms, 2),
            'generation_time_ms': round(self.generation_time_ms, 2),
            'total_time_ms': round(self.total_time_ms, 2),
            'retrieved_count': self.retrieved_count,
            'has_relevant_context': self.has_relevant_context,
        }

# ================== 文档清单管理 ==================
@dataclass
class DocumentRecord:
    """单文档记录"""
    file_path: str
    file_hash: str
    last_modified: float
    chunk_ids: list[str] = field(default_factory=list)
    chunk_count: int = 0
    file_type: str = ""
    indexed_at: str = ""


class DocumentManifest:
    """
    企业级文档清单管理器

    原理：
    - 维护已入库文档的哈希值和元数据映射，实现增量更新
    - 通过文件哈希(MD5) + 修改时间 双重检测文件变更
    - 支持原子性写入（先写临时文件再rename），防止写入中断导致清单损坏
    - 每次变更前自动备份，支持回滚
    """

    def __init__(self, manifest_path: str):
        self._path = manifest_path
        self._records: dict[str, DocumentRecord] = {}
        self._lock = None  # 线程锁，实际使用时由RAGSystem的锁控制

    def load(self) -> bool:
        """加载清单，失败时自动从备份恢复"""
        if not os.path.exists(self._path):
            return False
        try:
            with open(self._path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            self._records = {}
            for file_path, rec in data.get('documents', {}).items():
                self._records[file_path] = DocumentRecord(
                    file_path=rec.get('file_path', file_path),
                    file_hash=rec.get('file_hash', ''),
                    last_modified=rec.get('last_modified', 0),
                    chunk_ids=rec.get('chunk_ids', []),
                    chunk_count=rec.get('chunk_count', 0),
                    file_type=rec.get('file_type', ''),
                    indexed_at=rec.get('indexed_at', ''),
                )
            return True
        except (json.JSONDecodeError, KeyError):
            backup = self._path + '.backup'
            if os.path.exists(backup):
                try:
                    with open(backup, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    self._records = {}
                    for file_path, rec in data.get('documents', {}).items():
                        self._records[file_path] = DocumentRecord(**rec)
                    return True
                except Exception:
                    pass
            self._records = {}
            return False

    def save(self) -> None:
        """原子性保存清单（写临时文件 + rename）"""
        os.makedirs(os.path.dirname(self._path), exist_ok=True)
        data = {
            'version': '1.0',
            'updated_at': datetime.now().isoformat(),
            'total_documents': len(self._records),
            'total_chunks': sum(r.chunk_count for r in self._records.values()),
            'documents': {
                fp: {
                    'file_path': r.file_path,
                    'file_hash': r.file_hash,
                    'last_modified': r.last_modified,
                    'chunk_ids': r.chunk_ids,
                    'chunk_count': r.chunk_count,
                    'file_type': r.file_type,
                    'indexed_at': r.indexed_at,
                }
                for fp, r in self._records.items()
            }
        }
        tmp_path = self._path + '.tmp'
        with open(tmp_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

        if os.path.exists(self._path):
            try:
                os.replace(self._path, self._path + '.backup')
            except OSError:
                pass

        os.replace(tmp_path, self._path)

    def has_file(self, file_path: str) -> bool:
        return file_path in self._records

    def get_record(self, file_path: str) -> Optional[DocumentRecord]:
        return self._records.get(file_path)

    def add_record(self, record: DocumentRecord) -> None:
        self._records[record.file_path] = record

    def remove_record(self, file_path: str) -> Optional[DocumentRecord]:
        return self._records.pop(file_path, None)

    def get_all_paths(self) -> set[str]:
        return set(self._records.keys())

    def get_total_chunks(self) -> int:
        return sum(r.chunk_count for r in self._records.values())

    def get_total_documents(self) -> int:
        return len(self._records)


# ================== [OPT-04] 企业级评估器 ==================
class EnterpriseJudge:
    """
    企业级评估器 - 从关键词重叠升级为语义级评估

    原理说明：
    - 关键词重叠(Jaccard)忽略语义，2个句子可能0重叠但语义相同("DCT变换" vs "离散余弦变换")
    - 企业级评估需要: 语义相似度(embedding) + 事实匹配(cross-encoder) + 结构匹配(ROUGE)
    - 耗时优化: 优先使用轻量算法，仅在必要时使用LLM/cross-encoder
    """

    def __init__(self, embeddings=None, rerank_model=None):
        """
        :param embeddings: OllamaEmbeddings 实例，用于语义相似度计算
        :param rerank_model: CrossEncoder 实例，用于精细相关性判断
        """
        self._embeddings = embeddings
        self._rerank_model = rerank_model
        self._stopwords = self._init_stopwords()

    def _init_stopwords(self) -> set[str]:
        return {
            '的', '了', '是', '在', '我', '有', '和', '就', '不', '人', '都', '一',
            '一个', '上', '也', '很', '到', '说', '要', '去', '你', '会', '着',
            '没有', '看', '好', '这个', '那个', '什么', '怎么', '为什么', '哪里',
            '哪个', '如何', '请问', '一下', '吧', '吗', '呢', '啊', '哦', '嗯',
            '是的', '好的', '可以', '应该',
            'a', 'an', 'the', 'and', 'of', 'to', 'in', 'for', 'on', 'with',
            'by', 'at', 'is', 'are', 'was', 'were', 'be', 'been', 'being',
            'have', 'has', 'had',
        }

    # ------------------ 辅助方法 ------------------
    def _tokenize_keywords(self, text: str) -> set[str]:
        """提取文本中的非停用词集合"""
        words = set(jieba.cut(text))
        return {w for w in words if len(w) > 1 and w not in self._stopwords}

    def _split_sentences(self, text: str) -> list[str]:
        """智能分句，支持中英文标点"""
        if not text:
            return []
        raw = re.split(r'[。！？!?\n；;]+', text)
        return [s.strip() for s in raw if len(s.strip()) >= 4]

    def _extract_key_facts(self, text: str) -> list[str]:
        """
        从文本中提取关键事实点（以句子为单位）
        原理：评估召回率时，不是看词覆盖，而是看关键事实是否被检索到
        """
        sentences = self._split_sentences(text)
        if len(sentences) <= 3:
            return sentences
        # 按长度和位置加权选择关键句
        scored = []
        for i, sent in enumerate(sentences):
            keywords = self._tokenize_keywords(sent)
            score = len(keywords) * (1.0 - 0.3 * (i / max(len(sentences), 1)))  # 位置权重
            scored.append((sent, score))
        scored.sort(key=lambda x: x[1], reverse=True)
        # 取前一半或至少3句
        n = max(3, len(sentences) // 2)
        return [s[0] for s in scored[:n]]

    def _cosine_similarity(self, vec1, vec2) -> float:
        """计算向量余弦相似度"""
        if vec1 is None or vec2 is None:
            return 0.0
        dot = np.dot(vec1, vec2)
        norm = np.linalg.norm(vec1) * np.linalg.norm(vec2)
        return float(dot / norm) if norm > 0 else 0.0

    def _embed_text(self, text: str):
        """获取文本的embedding向量，带异常保护；同文本复用缓存避免重复调用"""
        if self._embeddings is None or not text or not text.strip():
            return None
        cache_key = text[:2000]
        if not hasattr(self, '_embed_cache'):
            self._embed_cache = {}
        if cache_key in self._embed_cache:
            return self._embed_cache[cache_key]
        try:
            vec = np.array(self._embeddings.embed_query(text))
            if len(self._embed_cache) < 64:
                self._embed_cache[cache_key] = vec
            return vec
        except Exception:
            return None

    def _compute_rouge_l(self, reference: str, candidate: str) -> float:
        """
        计算 ROUGE-L (Longest Common Subsequence) F1
        原理：LCS衡量答案与标准答案的结构相似度，比Jaccard更能反映内容正确性
        """
        if not reference or not candidate:
            return 0.0
        # 使用字符级别的LCS，对中文更友好
        ref_chars = list(reference)
        cand_chars = list(candidate)
        m, n = len(ref_chars), len(cand_chars)
        if m == 0 or n == 0:
            return 0.0
        # 动态规划求LCS长度
        dp = [[0] * (n + 1) for _ in range(m + 1)]
        for i in range(1, m + 1):
            for j in range(1, n + 1):
                if ref_chars[i - 1] == cand_chars[j - 1]:
                    dp[i][j] = dp[i - 1][j - 1] + 1
                else:
                    dp[i][j] = max(dp[i - 1][j], dp[i][j - 1])
        lcs_len = dp[m][n]
        recall = lcs_len / m if m > 0 else 0
        precision = lcs_len / n if n > 0 else 0
        if recall + precision == 0:
            return 0.0
        return 2 * recall * precision / (recall + precision)

    # ==================== 四大核心指标 ====================

    def compute_faithfulness(self, question: str, answer: str, contexts: list[str]) -> float:
        """
        [OPT-04] 企业级忠实度计算

        原理：
        - 原版：答案逐句与context做词重叠，30%即认为"支持"，阈值过低，导致虚假高分。
        - 优化：三阶段混合策略：
          Step1 (快速): 句子级关键词Jaccard预筛(阈值50%)，快速排除明显不支持的句子
          Step2 (中等): 对边界case用embedding语义相似度验证(阈值0.55)
          Step3 (精细): 当有cross-encoder时做深度校验

        最终忠实度 = 被支持的句子数 / 总句子数
        额外输出 hallucination_rate = 1 - faithfulness（幻觉率）
        """
        answer_sentences = self._split_sentences(answer)
        if not answer_sentences:
            return 1.0

        all_context = " ".join(contexts)
        context_words = self._tokenize_keywords(all_context)
        ctx_vec = self._embed_text(all_context[:1000])  # 上下文向量只算一次

        supported = 0
        for sent in answer_sentences:
            sent_words = self._tokenize_keywords(sent)
            if not sent_words:
                supported += 1
                continue

            # Step1: 关键词Jaccard（阈值从0.3提升到0.5）
            intersection = len(sent_words & context_words)
            jaccard = intersection / len(sent_words) if sent_words else 0

            if jaccard >= 0.5:  # [OPT-04] 从0.3提升到0.5
                supported += 1
            elif jaccard >= 0.2:
                # Step2: 边界case用语义相似度二次校验
                sent_vec = self._embed_text(sent)
                semantic_score = self._cosine_similarity(sent_vec, ctx_vec)
                if semantic_score >= 0.55:
                    supported += 1
                elif self._rerank_model is not None and jaccard >= 0.1:
                    # Step3: CrossEncoder精细校验（最慢但最准）
                    try:
                        pair_score = self._rerank_model.predict([[sent, all_context[:1500]]])[0]
                        if pair_score > 0.3:
                            supported += 1
                    except Exception:
                        pass
            # jaccard < 0.2: 明确不支持，不计数

        return supported / len(answer_sentences)

    def compute_context_recall(self, ground_truth: str, contexts: list[str]) -> float:
        """
        [OPT-04] 企业级上下文召回率

        原理：
        - 原版：ground_truth词集合 vs context词集合的Jaccard
          问题：Jaccard不区分关键词重要程度，"的"和"DCT"权重一样。
        - 优化：关键事实覆盖法
          Step1: 从ground_truth中提取关键事实点(按句子+位置权重)
          Step2: 每个关键事实点与context计算语义相似度
          Step3: 相似度>阈值的视为"已检索到"

        影响：能精确衡量"用户想知道的那些关键信息，检索系统是否找到了"
        """
        all_context = " ".join(contexts)
        if not ground_truth or not all_context:
            return 0.0
        if len(ground_truth.strip()) < 10:
            return 1.0

        # Step1: 提取关键事实
        key_facts = self._extract_key_facts(ground_truth)
        if not key_facts:
            gt_words = self._tokenize_keywords(ground_truth)
            ctx_words = self._tokenize_keywords(all_context)
            covered = len(gt_words & ctx_words)
            total = len(gt_words)
            return covered / total if total > 0 else 1.0

        # Step2: 每个关键事实与context做语义匹配
        ctx_vec = self._embed_text(all_context[:1500])  # 上下文向量只算一次
        ctx_words = self._tokenize_keywords(all_context)
        covered_facts = 0
        for fact in key_facts:
            fact_vec = self._embed_text(fact)
            sim = self._cosine_similarity(fact_vec, ctx_vec)

            if sim >= 0.45:  # 语义阈值：中等相关即可认为检索到
                covered_facts += 1
            else:
                # 快速词重叠作为兜底
                f_words = self._tokenize_keywords(fact)
                word_overlap = len(f_words & ctx_words) / max(len(f_words), 1)
                if word_overlap >= 0.4:
                    covered_facts += 1

        return covered_facts / len(key_facts)

    def compute_answer_accuracy(self, answer: str, ground_truth: str) -> float:
        """
        [OPT-04] 企业级答案准确性

        原理：
        - 原版：answer vs ground_truth的Jaccard
          问题："DCT(离散余弦变换)"和"DCT"词重叠低但语义一致
        - 优化：三维度加权融合
          (1) 语义相似度 (embedding cos) - 权重0.4
          (2) 关键事实匹配 (事实覆盖) - 权重0.35
          (3) ROUGE-L (结构匹配) - 权重0.25

        三维度互补：
        - 语义相似度处理"说法不同但意思相同"
        - 事实匹配确保核心信息点不遗漏
        - ROUGE-L捕捉答案结构的完整度
        """
        if not answer or not ground_truth:
            return 0.0
        if len(ground_truth.strip()) < 5:
            return 1.0

        # (1) 语义相似度
        ans_vec = self._embed_text(answer)
        gt_vec = self._embed_text(ground_truth)
        semantic_sim = self._cosine_similarity(ans_vec, gt_vec)

        # (2) 关键事实匹配：ground_truth中的关键事实是否在答案中出现
        key_facts = self._extract_key_facts(ground_truth)
        a_words = self._tokenize_keywords(answer)
        matched = 0
        for fact in key_facts:
            if ans_vec is not None:
                f_vec = self._embed_text(fact)
                if f_vec is not None:
                    sim = self._cosine_similarity(f_vec, ans_vec)
                    if sim >= 0.4:
                        matched += 1
                        continue
            # 快速词重叠兜底
            f_words = self._tokenize_keywords(fact)
            overlap = len(f_words & a_words) / max(len(f_words), 1)
            if overlap >= 0.35:
                matched += 1
        fact_match = matched / len(key_facts) if key_facts else semantic_sim

        # (3) ROUGE-L
        rouge = self._compute_rouge_l(ground_truth, answer)

        # 加权融合
        accuracy = 0.40 * semantic_sim + 0.35 * fact_match + 0.25 * rouge
        return min(1.0, accuracy)

    def compute_answer_relevance(self, answer: str, question: str) -> float:
        """
        [OPT-04] 企业级答案相关性

        原理：
        - 原版：answer vs question词重叠
          问题：答案包含很多技术术语但完全未回答用户问题也能得高分
        - 优化：双维度检查
          (1) 语义相关性：answer和question的embedding余弦相似度
          (2) 问题意图匹配：答案是否包含问题的核心疑问词和关键实体

        两者的几何平均确保：
        - 答案既在语义上和问题相关(不跑题)
        - 又确实在尝试回答用户问的问题(不答非所问)
        """
        if not answer or not question:
            return 0.0

        # (1) 语义相关性
        ans_vec = self._embed_text(answer)
        q_vec = self._embed_text(question)
        semantic_rel = self._cosine_similarity(ans_vec, q_vec)

        # (2) 问题意图匹配：提取问题中的核心实体和疑问
        q_keywords = self._tokenize_keywords(question)
        a_keywords = self._tokenize_keywords(answer)
        if not q_keywords:
            return semantic_rel

        # 核心实体覆盖：问题中提到的名词/关键词在答案中出现的比例
        entity_coverage = len(q_keywords & a_keywords) / len(q_keywords)

        # 几何平均：两个维度都必须高才算真正相关
        relevance = (semantic_rel * entity_coverage) ** 0.5
        return min(1.0, relevance)

class RAGEvaluator:
    """RAG系统评估器"""
    def __init__(self, rag_system):
        self.rag_system = rag_system
        # [OPT-04] 使用企业级Judge替换SimpleJudge，传入模型引用
        embeddings = rag_system.embeddings if hasattr(rag_system, 'embeddings') else None
        rerank_model = rag_system.rerank_model if hasattr(rag_system, 'rerank_model') else None
        self.judge = EnterpriseJudge(embeddings=embeddings, rerank_model=rerank_model)
        self.results: List[EvalResult] = []

    def load_testset(self, csv_path: str) -> pd.DataFrame:
        """加载测试集并进行数据质量检查"""
        possible_paths = [
            csv_path,
            f"./{csv_path}",
            f"../{csv_path}",
            f"../../{csv_path}",
            f"./RAG_files/{csv_path}",
        ]
        found = csv_path
        for path in possible_paths:
            if os.path.exists(path):
                found = path
                break
        else:
            raise FileNotFoundError(f"找不到测试集文件: {csv_path}")

        df = pd.read_csv(found, encoding='utf-8')
        # 数据质量检查
        required_cols = ['question', 'ground_truth']
        missing = [c for c in required_cols if c not in df.columns]
        if missing:
            raise ValueError(f"测试集缺少必要列: {missing}")

        print(f"加载测试集: {os.path.basename(found)}, 共{len(df)}条")
        return df

    @retry_on_failure(max_retries=1, delay=2.0)
    def _retrieve_safe(self, question: str) -> List:
        """带重试保护的检索调用"""
        return self.rag_system._retrieve_with_rerank(question)

    @retry_on_failure(max_retries=1, delay=2.0)
    def _generate_safe(self, question: str, contexts: list = None):
        """带重试保护的生成调用；传入contexts时跳过重复检索"""
        if self.rag_system.rag_chain is None:
            self.rag_system._build_rag_chain()
        if contexts is not None:
            return self.rag_system._generate_answer(question, contexts)
        return self.rag_system.rag_chain.invoke({"input": question})

    def evaluate_single(self, row: pd.Series) -> EvalResult:
        """评估单个问题"""
        question = str(row.get('question', '') or '')
        ground_truth = str(row.get('ground_truth', '') or '')
        question_type = str(row.get('question_type', 'unknown') or 'unknown')
        difficulty = str(row.get('difficulty', 'unknown') or 'unknown')

        result = EvalResult(
            question=question,
            ground_truth=ground_truth,
            retrieved_answer='',
            retrieved_contexts=[],
            question_type=question_type,
            difficulty=difficulty,
        )

        total_start = time.time()
        retrieval_start = time.time()

        try:
            # 检索阶段
            contexts = self._retrieve_safe(question)
            result.retrieved_contexts = [
                doc.page_content if hasattr(doc, 'page_content') else str(doc)
                for doc in (contexts or [])
            ]
            result.retrieved_count = len(result.retrieved_contexts)
            result.retrieval_time_ms = (time.time() - retrieval_start) * 1000

            # 检查检索相关性
            if ground_truth and len(ground_truth) > 5:
                all_context = " ".join(result.retrieved_contexts)
                gt_words = self.judge._tokenize_keywords(ground_truth)
                ctx_words = self.judge._tokenize_keywords(all_context)
                if gt_words:
                    overlap = len(gt_words & ctx_words) / len(gt_words)
                    result.has_relevant_context = overlap > 0.15

            # 生成阶段（复用已检索context，避免rag_chain内二次检索）
            generation_start = time.time()
            response = self._generate_safe(question, contexts=contexts)
            result.retrieved_answer = response if isinstance(response, str) else str(response)
            result.generation_time_ms = (time.time() - generation_start) * 1000
            result.total_time_ms = (time.time() - total_start) * 1000

        except Exception as e:
            result.total_time_ms = (time.time() - total_start) * 1000
            return result

        # [OPT-04] 企业级指标计算
        if ground_truth and len(ground_truth) > 0:
            try:
                result.answer_accuracy = self.judge.compute_answer_accuracy(
                    result.retrieved_answer, ground_truth
                )
                result.answer_relevance = self.judge.compute_answer_relevance(
                    result.retrieved_answer, question
                )
                result.context_recall = self.judge.compute_context_recall(
                    ground_truth, result.retrieved_contexts
                )
                result.faithfulness = self.judge.compute_faithfulness(
                    question, result.retrieved_answer, result.retrieved_contexts
                )

                # 子维度指标
                if result.faithfulness is not None:
                    result.hallucination_rate = max(0.0, 1.0 - result.faithfulness)

            except Exception:
                pass

        return result

    def run_evaluation(self, testset_path: str, limit: int = None,
                       save_results: bool = True) -> dict:
        """运行完整评估"""
        df = self.load_testset(testset_path)
        valid_df = df[
            df['ground_truth'].notna() &
            (df['ground_truth'].astype(str).str.strip() != '') &
            (df['question'].notna()) &
            (df['question'].astype(str).str.strip() != '')
        ]
        print(f"有效测试用例: {len(valid_df)}/{len(df)}")

        if limit:
            valid_df = valid_df.head(limit)
            print(f"限制评估数量: {limit}")

        type_counts = valid_df['question_type'].value_counts()
        print("测试用例分布:")
        for qtype, count in type_counts.items():
            print(f"   {qtype}: {count}")

        print("开始评估...")
        self.results = []

        for idx, (_, row) in enumerate(tqdm(
            valid_df.iterrows(), total=len(valid_df), desc="评估进度"
        )):
            result = self.evaluate_single(row)
            self.results.append(result)

            if (idx + 1) % 5 == 0 and self.results:
                recent = self.results[-5:]
                avg_acc = np.mean([r.answer_accuracy for r in recent])
                avg_faith = np.mean([r.faithfulness for r in recent])

        summary = self._compute_summary()
        self._print_summary(summary)

        if save_results:
            self._save_results()

        return summary

    def _compute_summary(self) -> dict:
        """计算汇总指标 - 扩展维度"""
        if not self.results:
            return {}

        def safe_mean(values):
            valid = [v for v in values if v is not None]
            return float(np.mean(valid)) if valid else 0.0

        metrics = {
            'total_questions': len(self.results),
            'avg_answer_accuracy': safe_mean([r.answer_accuracy for r in self.results]),
            'avg_answer_relevance': safe_mean([r.answer_relevance for r in self.results]),
            'avg_faithfulness': safe_mean([r.faithfulness for r in self.results]),
            'avg_context_recall': safe_mean([r.context_recall for r in self.results]),
            'avg_hallucination_rate': safe_mean([r.hallucination_rate for r in self.results]),
            'avg_retrieval_time_ms': safe_mean([r.retrieval_time_ms for r in self.results]),
            'avg_generation_time_ms': safe_mean([r.generation_time_ms for r in self.results]),
            'avg_total_time_ms': safe_mean([r.total_time_ms for r in self.results]),
            'avg_retrieved_count': safe_mean([r.retrieved_count for r in self.results]),
            'has_relevant_context_rate': safe_mean([
                1 if r.has_relevant_context else 0 for r in self.results
            ]),
            'success_rate': safe_mean([
                1 if r.answer_accuracy > 0.3 else 0 for r in self.results
            ]),
            'excellent_rate': safe_mean([
                1 if r.answer_accuracy > 0.6 else 0 for r in self.results
            ]),
        }

        # 按难度分组
        difficulties = set(r.difficulty for r in self.results if r.difficulty)
        metrics['by_difficulty'] = {}
        for diff in difficulties:
            dr = [r for r in self.results if r.difficulty == diff]
            if dr:
                metrics['by_difficulty'][diff] = {
                    'count': len(dr),
                    'avg_accuracy': safe_mean([r.answer_accuracy for r in dr]),
                    'avg_faithfulness': safe_mean([r.faithfulness for r in dr]),
                    'avg_recall': safe_mean([r.context_recall for r in dr]),
                }

        # 按问题类型分组
        question_types = set(r.question_type for r in self.results if r.question_type)
        metrics['by_question_type'] = {}
        for qtype in question_types:
            tr = [r for r in self.results if r.question_type == qtype]
            if tr:
                metrics['by_question_type'][qtype] = {
                    'count': len(tr),
                    'avg_accuracy': safe_mean([r.answer_accuracy for r in tr]),
                    'avg_faithfulness': safe_mean([r.faithfulness for r in tr]),
                    'avg_recall': safe_mean([r.context_recall for r in tr]),
                }

        return metrics

    def _print_summary(self, summary: dict):
        """打印汇总报告"""
        if not summary:
            return
        print("\n" + "=" * 60)
        print("  评估结果汇总")
        print("=" * 60)
        print(f"  总问题数:           {summary.get('total_questions', 0)}")
        print(f"  答案准确性(avg):    {summary.get('avg_answer_accuracy', 0):.4f}")
        print(f"  答案相关性(avg):    {summary.get('avg_answer_relevance', 0):.4f}")
        print(f"  忠实度(avg):        {summary.get('avg_faithfulness', 0):.4f}")
        print(f"  上下文召回率(avg):  {summary.get('avg_context_recall', 0):.4f}")
        print(f"  幻觉率(avg):        {summary.get('avg_hallucination_rate', 0):.4f}")
        print(f"  相关上下文命中率:   {summary.get('has_relevant_context_rate', 0):.2%}")
        print(f"  成功率(acc>0.3):    {summary.get('success_rate', 0):.2%}")
        print(f"  优秀率(acc>0.6):    {summary.get('excellent_rate', 0):.2%}")
        print(f"  平均检索耗时:       {summary.get('avg_retrieval_time_ms', 0):.0f}ms")
        print(f"  平均生成耗时:       {summary.get('avg_generation_time_ms', 0):.0f}ms")
        print(f"  平均总耗时:         {summary.get('avg_total_time_ms', 0):.0f}ms")
        print("=" * 60)

    def _save_results(self):
        """保存评估结果"""
        os.makedirs(EVAL_RESULTS_DIR, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        results_df = pd.DataFrame([r.to_dict() for r in self.results])
        results_path = os.path.join(EVAL_RESULTS_DIR, f"eval_results_{timestamp}.csv")
        results_df.to_csv(results_path, index=False, encoding='utf-8')
        print(f"\n详细结果: {results_path}")

        summary = self._compute_summary()
        summary_path = os.path.join(EVAL_RESULTS_DIR, f"eval_summary_{timestamp}.json")
        with open(summary_path, 'w', encoding='utf-8') as f:
            json.dump(summary, f, ensure_ascii=False, indent=2)
        print(f"汇总指标: {summary_path}")

class LocalRAGSystem:
    def __init__(self):
        #懒加载配置
        self._embeddings = None
        self._chat_model = None
        self._multi_query_model = None
        self._compression_model = None
        self._rerank_model = None
        self.vector_store = None
        self.rag_chain = None
        self.chunks: List[LangchainDocument] = []
        self._bm25_index = None
        self._keyword_index = None
        self._stopwords: Set[str] = set()
        self._document_chain = None
        self._rag_prompt = None
        self._init_stopwords()
        self._manifest = DocumentManifest(MANIFEST_PATH)
        self._mutation_lock = threading.Lock()
    
    @property
    def embeddings(self):
        if self._embeddings is None:
            self._embeddings = OllamaEmbeddings(
            model=EMBED_MODEL
            )
        return self._embeddings
    
    @property
    def chat_model(self):
        if self._chat_model is None:
            self._chat_model = self._create_chat_model(streaming=CHAT_STREAMING)
        return self._chat_model

    def _create_chat_model(self, streaming: bool = True):
        """创建 ChatOllama 实例；streaming 仅交互模式需要"""
        kwargs = dict(
            model=CHAT_MODEL,
            temperature=0.3,
            num_threads=12,
            num_ctx=CHAT_NUM_CTX,
            #num_predict=CHAT_NUM_PREDICT,
            keep_alive=OLLAMA_KEEP_ALIVE,
            top_p=0.7,
            repeat_penalty=1.1,
            streaming=streaming,
        )
        if streaming:
            kwargs["callbacks"] = [StreamingStdOutCallbackHandler()]
        return ChatOllama(**kwargs)
    
    @property
    def rerank_model(self):
        if self._rerank_model is None:
            # [OPT-11] 自动检测设备，优先GPU
            try:
                import torch
                device = "cuda" if torch.cuda.is_available() else "cpu"
            except ImportError:
                device = "cpu"
            self._rerank_model = CrossEncoder(
                "./models/bge-reranker-base", device=device
            )
        return self._rerank_model
    
    @property
    def multi_query_model(self):
        """[OPT-02修复] 改为property，原版是普通方法且逻辑错误"""
        if self._multi_query_model is None and ENABLE_MULTI_QUERY:
            self._multi_query_model = ChatOllama(
                model=MULTI_QUERY_MODEL,
                temperature=0.1,
                num_threads=12,
                num_ctx=4096,
                streaming=False,
            )
        return self._multi_query_model
    @property
    def compression_model(self):
        if self._compression_model is None and ENABLE_DOC_COMPRESSION:
            self._compression_model = ChatOllama(
                model=COMPRESSION_MODEL,
                temperature=COMPRESSION_TEMPERATURE,
                num_threads=12,
                num_ctx=2048,
                streaming=False,
            )
        return self._compression_model

    def _init_stopwords(self):
        self._stopwords = {
            '的', '了', '是', '在', '我', '有', '和', '就', '不', '人', '都', '一',
            '一个', '上', '也', '很', '到', '说', '要', '去', '你', '会', '着',
            '没有', '看', '好', '这个', '那个', '什么', '怎么', '为什么', '哪里',
            '哪个', '如何', '请问', '一下',
            '吧', '吗', '呢', '啊', '哦', '嗯', '哈哈', '是的', '好的', '可以', '应该',
            'a', 'an', 'the', 'and', 'of', 'to', 'in', 'for', 'on', 'with',
            'by', 'at', 'is', 'are', 'was', 'were', 'be', 'been', 'being',
            'have', 'has', 'had',
        }
    def _load_docx_with_heading(self, docx_path: str) -> LangchainDocument:
        """
        [OPT-13] 增强版Word文档加载：兼容.doc/.docx，更健壮的标题检测
        原理：原版仅通过style name正则匹配，对自定义样式文档无法识别标题。
              增加对内置Heading样式ID的检测，覆盖更多文档类型。
        """
        try:
            word_doc = Document(docx_path)
        except Exception:
            print("=================")
            return LangchainDocument(
                page_content=f"[加载失败] {docx_path}",
                metadata={"source": docx_path, "load_error": True}
            )

        markdown_lines = []
        for para in word_doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = (para.style.name or "").lower() if para.style else ""
            level = 0

            # 方法1：样式名称正则匹配 (标题 1, Heading 1, head 1)
            match = re.search(r'(?:head(?:ing)?|标题)\s*(\d+)', style_name, re.IGNORECASE)
            if match:
                level = int(match.group(1))
            else:
                # [OPT-13] 方法2：通过样式ID判断内置标题
                style_id = (para.style.style_id or "").lower() if para.style else ""
                match_id = re.search(r'heading(\d+)', style_id)
                if match_id:
                    level = int(match_id.group(1))

            if 1 <= level <= 3:
                markdown_lines.append(f"{'#' * level} {text}")
            else:
                markdown_lines.append(text)

        full_text = "\n\n".join(markdown_lines)
        return LangchainDocument(
            page_content=full_text,
            metadata={"source": docx_path, "file_type": "docx"}
        )

    def _load_pdf_with_fallback(self, pdf_path: str) -> List[LangchainDocument]:
        """
        [OPT-05] 带容错的PDF加载：先尝试完整加载，失败则逐页加载。
        原理：PyPDFLoader的load()一次性加载所有页，遇到损坏页会整体失败；
              逐页加载可以跳过损坏页，保留有效页的数据。
        """
        try:
            loader = PyPDFLoader(str(pdf_path))
            docs = loader.load()
            if docs:
                for i, doc in enumerate(docs):
                    doc.metadata["page"] = i + 1
                    doc.metadata["source"] = str(pdf_path)
                    doc.metadata["file_type"] = "pdf"
                return docs
        except Exception:
            pass

        # 逐页加载容错
        docs = []
        try:
            loader = PyPDFLoader(str(pdf_path))
            for page_num in range(1000):  # 安全上限
                try:
                    page_docs = loader.load()
                    for doc in page_docs:
                        if doc.metadata.get("page", -1) == page_num:
                            docs.append(doc)
                            break
                except Exception:
                    break
        except Exception:
            pass

        if not docs:
            docs = [LangchainDocument(
                page_content=f"[PDF加载失败] {os.path.basename(pdf_path)}",
                metadata={"source": str(pdf_path), "load_error": True, "file_type": "pdf"}
            )]

        return docs
    def _load_documents(self) -> List[LangchainDocument]:
        """
        [OPT-06] 并行加载目录下所有文档（txt/pdf/docx/doc）
        原理：文档加载是I/O密集型操作，并行加载可大幅减少启动时间。
        """
        docs: List[LangchainDocument] = []
        docs_path = Path(DOCS_DIR)

        if not docs_path.exists():
            docs_path.mkdir(parents=True)
            print(f"已创建文档目录: {DOCS_DIR}，请放入文档后重新运行")
            return []

        # 收集所有文件
        txt_files = list(docs_path.glob('*.txt'))
        pdf_files = list(docs_path.glob('*.pdf'))
        docx_files = list(docs_path.glob('*.docx'))
        doc_files = list(docs_path.glob('*.doc'))
        all_files = txt_files + pdf_files + docx_files + doc_files

        if not all_files:
            print(f"目录 {DOCS_DIR} 下未找到文档")
            return []

        print(f"发现 {len(all_files)} 个文档文件 (txt:{len(txt_files)} pdf:{len(pdf_files)} docx/doc:{len(docx_files)+len(doc_files)})")

        # [OPT-06] 并行加载
        def load_single_file(file_path: Path):
            ext = file_path.suffix.lower()
            try:
                if ext == '.txt':
                    loader = TextLoader(str(file_path), encoding='utf-8')
                    loaded = loader.load()
                    return loaded

                elif ext == '.pdf':
                    return self._load_pdf_with_fallback(str(file_path))

                elif ext in ('.docx', '.doc'):
                    doc = self._load_docx_with_heading(str(file_path))
                    return [doc]

            except Exception as e:
                return [LangchainDocument(
                    page_content=f"[加载失败] {file_path.name}: {str(e)}",
                    metadata={"source": str(file_path), "load_error": True}
                )]

        results = []
        with ThreadPoolExecutor(max_workers=min(MAX_WORKERS, len(all_files))) as executor:
            futures = {executor.submit(load_single_file, f): f for f in all_files}
            for future in as_completed(futures):
                file_path = futures[future]
                try:
                    loaded = future.result(timeout=60)
                    if loaded:
                        results.extend(loaded)
                except Exception as e:
                    pass

        docs = results
        print(f"文档加载完成，共 {len(docs)} 个文档对象")
        return docs

        # for suffix in suffix_list:
        #     file_path = list(docs_path.glob(f'*{suffix}'))
        #     for file in file_path:
        #         print(f"正在加载{suffix}文件，{file.name}")
        #         try:
        #             #使用UnstructuredLoader加载文档，必须安装对应的依赖包
        #             loader = UnstructuredLoader(
        #                 str(file),
        #                 mode = "elements", #开启精细元素解析
        #                 strategy = "hi_res", #高精度解析
        #                 extract_images_in_pdf=True,#开启，PDF 里图片的提取 + OCR 文字识别
        #                 pdf_infer_table_structure=True, #开启，自动识别PDF里面的表格，把表格解析成可读文本
        #                 extract_image_block_types=["Image"], #自动提取图片里的文字OCR识别
        #                 chunking_strategy="by_title",  # 自动按章节分块
        #                 include_page_breaks=False,        # 不保留分页符
        #                 languages=["chi_sim", "eng"],     # 支持中英文
        #             )
        #             doc = loader.load()
        #             docs.extend(doc)
        #             print(f"已加载{suffix}文件，{file.name}")

        #         except Exception as e:
        #             print(f"加载{suffix}文件{file.name}失败，{str(e)}")

      

    def _split_documents(self, docs: List[LangchainDocument]) -> List[LangchainDocument]:
        """
        [OPT-01] 关键修复：对所有文档类型执行文本切分。

        原版bug：if/elif逻辑导致只有.docx被切分，.txt和.pdf加载后直接丢弃。
        修复：统一使用RecursiveCharacterTextSplitter对所有文档切分，
              .docx额外经过MarkdownHeaderTextSplitter保留章节结构。

        原理：文本切分是向量化和检索的基础，未经切分的文档无法进入向量库。
        """
        final_chunks: List[LangchainDocument] = []
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            separators=['\n\n', '\n', '。 ', '. ', '！', '？', '!', '?', ' ', ''],
        )

        for doc in docs:
            file_name = os.path.basename(doc.metadata.get('source', ''))
            content = doc.page_content or ''
            if len(content.strip()) < 5:
                continue

            # [OPT-01] .docx 额外经过标题分层切分
            if file_name.endswith(('.docx', '.doc')):
                try:
                    headers_to_split_on = [
                        ("#", "H1"),
                        ("##", "H2"),
                        ("###", "H3"),
                    ]
                    md_splitter = MarkdownHeaderTextSplitter(
                        headers_to_split_on=headers_to_split_on
                    )
                    header_splits = md_splitter.split_text(content)
                    
                    chunks = splitter.split_documents(header_splits)
                    print(f"已切分{file_name}，共 {len(chunks)} 个文本块")
                except Exception:
                    chunks = splitter.split_documents([doc])
            else:
                # [OPT-01] txt/pdf 统一切分
                chunks = splitter.split_documents([doc])

            # 补充源文件元数据
            for chunk in chunks:
                if "source" not in chunk.metadata:
                    chunk.metadata["source"] = doc.metadata.get("source", file_name)
                if "file_type" not in chunk.metadata:
                    chunk.metadata["file_type"] = doc.metadata.get("file_type", "unknown")

            final_chunks.extend(chunks)

        print(f"文本切分完成，共 {len(final_chunks)} 个文本块")
        return final_chunks

    def build_or_load_db(self) -> None:
        """构建或加载向量数据库"""
        if os.path.exists(CHROMA_DB_DIR) and os.listdir(CHROMA_DB_DIR):
            print(f"加载已有向量库: {CHROMA_DB_DIR}")
            try:
                self.vector_store = Chroma(
                    persist_directory=CHROMA_DB_DIR,
                    embedding_function=self.embeddings,
                    collection_name='knowledge_base',
                )
                count = self.vector_store._collection.count()
                print(f"向量库加载完成，共 {count} 个向量")

                # 恢复chunks列表
                all_docs_data = self.vector_store.get()
                if all_docs_data and 'documents' in all_docs_data:
                    self.chunks = []
                    for doc_id, text, metadata in zip(
                        all_docs_data.get('ids', []),
                        all_docs_data.get('documents', []),
                        all_docs_data.get('metadatas', []),
                    ):
                        doc = LangchainDocument(
                            page_content=text,
                            metadata=metadata or {}
                        )
                        self.chunks.append(doc)

                self._manifest.load()
                if self._manifest.get_total_documents() == 0 and self.chunks:
                    self._rebuild_manifest_from_chunks()
            except Exception as e:
                print(f"加载向量库失败: {e}，将重新构建")
                self._build_new_db()
        else:
            self._build_new_db()

        # 构建检索索引
        if self.chunks:
            self._build_bm25_index()
            self._build_keyword_index()

        # [OPT-15] 后台预加载 embedding + rerank（chat 模型在 warmup_models 中加载）
        if PRELOAD_MODELS:
            self._preload_retrieval_models()

    def _preload_retrieval_models(self):
        """后台预加载检索侧模型，不阻塞主流程"""
        def _load():
            try:
                self.embeddings.embed_query("warmup")
                if ENABLE_RERANK and not FAST_MODE:
                    self.rerank_model.predict([["warmup", "warmup"]], batch_size=1)
            except Exception:
                pass

        ThreadPoolExecutor(max_workers=1).submit(_load)

    def warmup_models(self, preload_chat: bool = True):
        """
        [OPT-15] 显式预热所有模型。
        首次加载 qwen2.5:7b 可能需要 60-180s，应放在启动阶段而非首次问答时。
        """
        print("正在预热模型（首次加载较慢，请稍候）...")
        t_total = time.time()

        t = time.time()
        try:
            self.embeddings.embed_query("warmup")
            print(f"  Embedding 就绪  {time.time() - t:.1f}s")
        except Exception as e:
            print(f"  Embedding 失败: {e}")

        if ENABLE_RERANK and not FAST_MODE:
            t = time.time()
            try:
                self.rerank_model.predict([["warmup", "warmup"]], batch_size=1)
                print(f"  Rerank 就绪      {time.time() - t:.1f}s")
            except Exception as e:
                print(f"  Rerank 失败: {e}")

        if preload_chat:
            t = time.time()
            try:
                self.chat_model.invoke("你好")
                print(f"  Chat LLM 就绪    {time.time() - t:.1f}s")
            except Exception as e:
                print(f"  Chat LLM 失败: {e}")

        print(f"模型预热完成，总耗时 {time.time() - t_total:.1f}s\n")

    def _build_new_db(self):
        """构建新的向量库"""
        print("首次运行，开始构建向量库...")
        docs = self._load_documents()
        if not docs:
            print("未找到文档，无法构建向量库")
            return

        self.chunks = self._split_documents(docs)
        if not self.chunks:
            print("文档切分后无有效内容")
            return

        for i, chunk in enumerate(self.chunks):
            chunk.metadata["chunk_index"] = i

        print(f"正在向量化 {len(self.chunks)} 个文本块...")
        try:
            self.vector_store = Chroma.from_documents(
                documents=self.chunks,
                persist_directory=CHROMA_DB_DIR,
                embedding=self.embeddings,
                collection_name='knowledge_base',
            )
            print(f"向量库构建完成，共 {len(self.chunks)} 个向量")
            self._rebuild_manifest_from_chunks()
        except Exception as e:
            print(f"向量库构建失败: {e}")
            raise
    # ==================== 索引构建 ====================
    def _build_bm25_index(self):
        """构建BM25索引"""
        tokenized_docs = []
        for d in self.chunks:
            tokens = [t for t in jieba.cut(d.page_content.strip())
                      if t.strip() and t not in self._stopwords]
            tokenized_docs.append(tokens)
        self._bm25_index = BM25Okapi(tokenized_docs)

    def _build_keyword_index(self):
        """构建倒排关键词索引"""
        self._keyword_index = {}
        for idx, doc in enumerate(self.chunks):
            words = set(jieba.cut(doc.page_content.strip()))
            words = {w for w in words if w not in self._stopwords and len(w) > 1}
            for word in words:
                if word not in self._keyword_index:
                    self._keyword_index[word] = set()
                self._keyword_index[word].add(idx)
    # ================== 文档压缩（LLM抽取式） ==================
    def _compress_document(self, document: LangchainDocument, query: str) -> LangchainDocument:
        """
        使用LLM对单个文档进行抽取式压缩
        提取与用户问题最相关的核心信息，去除冗余内容       
        :param document: 原始文档对象
        :param query: 用户问题
        :return: 压缩后的文档对象
        """
        if not ENABLE_DOC_COMPRESSION or self.compression_model is None:
            return document
        
        original_content = document.page_content
        original_length = len(original_content)
        # 如果文档已经很短，不需要压缩
        target_len = COMPRESSION_TARGET_LENGTH
        if COMPRESSION_RATIO < 1.0:
            target_len = int(original_length * COMPRESSION_RATIO)
        
        if original_length <= target_len:
            return document
        
        # 构建压缩提示词模板
        compression_prompt = ChatPromptTemplate.from_messages([
            ("system", """你是文档压缩助手。根据用户问题，从资料中提取最相关的核心信息。
压缩规则：
1. 严格保留原文，不添加原文没有的信息
2. 只提取与用户问题直接相关的内容
3. 保留重要的事实、数据、定义、结论
4. 去除重复、举例、过渡性语句
5. 控制输出长度约{target_len}字符
直接输出压缩后的文本。"""),
            ("human", """用户问题：{query}

原始资料：
{document_content}
""")
        ])
        try:
            # 执行压缩
            chain = compression_prompt | self.compression_model
            response = chain.invoke({
                "query": query,
                "document_content": original_content,
                "target_len": target_len
            })
            
            compressed = response.content.strip() if hasattr(response, 'content') else str(response).strip()
            
            # 如果压缩失败或压缩后为空，返回原文
            if not compressed or len(compressed) < 10:
                print(f"⚠️ 文档压缩失败，使用原文")
                return document
            
            compression_ratio = len(compressed) / original_length
            print(f"📦 文档压缩完成: {original_length} → {len(compressed)} 字符 (压缩率: {compression_ratio:.1%})")
            # 创建压缩后的新文档，保留原元数据并添加压缩信息
            compressed_doc = LangchainDocument(
                page_content=compressed,
                metadata={
                    **document.metadata,
                    "compressed": True,
                    "original_length": original_length,
                    "compressed_length": len(compressed)
                }
            )
            return compressed_doc
        except Exception as e:
            print(f"⚠️ 文档压缩出错: {e}，使用原文")
            return document
    # ================== 批量文档压缩 ==================
    def _compress_documents_batch(self, documents: list[LangchainDocument], query: str) -> list[LangchainDocument]:
        """
        批量压缩文档（可并行处理，这里先实现串行）
        
        :param documents: 原始文档列表
        :param query: 用户问题
        :return: 压缩后的文档列表
        """
        if not ENABLE_DOC_COMPRESSION or self.compression_model is None:
            return documents
        
        print(f"\n🗜️ 开始文档压缩 (共{len(documents)}个文档)...")
        tasks = []
        compressed_docs = []
        with ThreadPoolExecutor(max_workers = MAX_WORKERS) as executor:
            for i,doc in enumerate(documents):
                tasks.append(executor.submit(self._compress_document, doc, query))
            
            for task in as_completed(tasks):
                try:
                    compressed = task.result(timeout=30)
                    if compressed:
                        compressed_docs.append(compressed)
                except Exception as e:
                    print(f"⚠️ 文档异步压缩出错: {e}")
                    pass

        print(f"✅ 文档压缩完成\n")
        return compressed_docs

    # ================== 备用：句子级别过滤压缩 ==================
    def _compress_by_sentence_filtering(self, document: LangchainDocument, query: str, top_n_sentences: int = 5) -> LangchainDocument:
        """
        基于句子相关性过滤的压缩方法（不依赖LLM，速度快）
        使用简单的关键词匹配和位置权重选择最相关的句子        
        :param document: 原始文档对象
        :param query: 用户问题
        :param top_n_sentences: 保留的最相关句子数量
        :return: 压缩后的文档对象
        """
        original_content = document.page_content
        # 分句（支持中英文标点）
        sentences = re.split(r'[。！？!?.\n]+', original_content)
        sentences = [s.strip() for s in sentences if len(s.strip()) > 10]
        
        if len(sentences) <= top_n_sentences:
            return document
        
        # 提取查询关键词
        query_keywords = set(jieba.cut(query))
        query_keywords = {kw for kw in query_keywords if kw not in self._stopwords and len(kw) > 1}
        # 计算每个句子的相关性分数
        sentence_scores = []
        for i, sent in enumerate(sentences):
            # 关键词匹配得分
            sent_words = set(jieba.cut(sent))
            sent_words = {sw for sw in sent_words if sw not in self._stopwords and len(sw) > 1}
            keyword_match_count = len(query_keywords & sent_words)
            keyword_score = keyword_match_count / max(len(query_keywords), 1)
            
            # 位置得分（越靠前权重越高）
            position_score = 1.0 - (i / len(sentences)) * 0.5
            
            # 长度惩罚（过长的句子可能包含冗余）
            length_penalty = min(1.0, 100.0 / len(sent))
            
            total_score = keyword_score * 0.6 + position_score * 0.3 + length_penalty * 0.1
            sentence_scores.append((i, sent, total_score))
        # 选择得分最高的句子
        sentence_scores.sort(key=lambda x: x[2], reverse=True)
        selected_sentences = [s[1] for s in sentence_scores[:top_n_sentences]]
        
        # 按原始顺序重新排列
        selected_indices = set(s[0] for s in sentence_scores[:top_n_sentences])
        ordered_sentences = [sentences[i] for i in range(len(sentences)) if i in selected_indices]
        
        compressed_content = "。".join(ordered_sentences)
        
        print(f"✂️ 句子过滤压缩: {len(sentences)}句 → {len(ordered_sentences)}句")
        
        return LangchainDocument(
            page_content=compressed_content,
            metadata={
                **document.metadata,
                "compressed": True,
                "compression_method": "sentence_filtering"
            }
        )
        
    # ================== 智能压缩（自动选择方法） ==================
    def _smart_compress_documents(self, documents: list[LangchainDocument], query: str) -> list[LangchainDocument]:
        """
        智能文档压缩：根据文档长度和配置选择压缩策略
        - 使用LLM抽取式压缩（效果好但慢）
        - 或使用句子过滤压缩（快但效果略差）
        """
        if not ENABLE_DOC_COMPRESSION:
            return documents
        
        # 计算总长度
        total_length = sum(len(doc.page_content) for doc in documents)
        
        # 如果总长度已经不大，不需要压缩
        if total_length < COMPRESSION_TARGET_LENGTH * 2:
            print(f"📄 文档总长度({total_length}字符)较小，跳过压缩")
            return documents
        # 选择压缩方法
        # 这里默认使用LLM抽取式压缩，如果想用快速压缩可以切换
        use_llm_compression = True  # 可以改成配置项
        
        if use_llm_compression:
            return self._compress_documents_batch(documents, query)
        else:
            compressed = []
            for doc in documents:
                compressed.append(self._compress_by_sentence_filtering(doc, query))
            return compressed

    def _extract_keywords_llm(self, query: str) -> list[str]:
        """
        使用LLM提取查询关键词
        :param query: 用户原始问题
        :return: 关键词列表
        """
        keyword_prompt = ChatPromptTemplate.from_messages([
            ("system", """你是一个关键词提取专家。请从用户问题中提取最核心的关键词，用于检索相关文档。
规则：
1. 只提取名词、专有名词、核心动词，忽略停用词（的、了、是、在等）
2. 每个关键词应该是独立且有意义的检索词
3. 关键词数量控制在{max_keywords}个以内
4. 只输出关键词，用逗号分隔，不要输出其他内容
"""),
            ("human", "{query}")
        ])
        
        try:
            chain = keyword_prompt | self.chat_model
            response = chain.invoke({
                "query": query,
                "max_keywords": MAX_KEYWORDS
            })
            # 解析LLM返回的关键词
            keywords = [kw.strip() for kw in response.content.split('，') if kw.strip()]
            # 如果LLM返回的是英文逗号分隔
            if len(keywords) <= 1 and ',' in response.content:
                keywords = [kw.strip() for kw in response.content.split(',') if kw.strip()]
            print(f"🔑 LLM提取的关键词: {keywords}")
            return keywords[:MAX_KEYWORDS]
        except Exception as e:
            print(f"⚠️ LLM关键词提取失败，回退到规则提取: {e}")
            return self._extract_keywords_rule_based(query)

    def _extract_keywords_rule_based(self, query: str) -> list[str]:
        """
        基于规则的关键词提取（轻量级，不依赖LLM）
        使用jieba分词 + 词性过滤 + TF-IDF/词频统计       
        :param query: 用户原始问题
        :return: 关键词列表
        """
        import jieba.posseg as pseg
        
        # 1. 分词并标注词性
        words = pseg.cut(query)
        
        # 2. 词性过滤：保留名词(n)、动名词(vn)、动词(v)、专有名词(nr, ns, nt)
        allowed_pos = ['n', 'vn', 'v', 'nr', 'ns', 'nt', 'eng']  # eng是英文单词
        
        candidate_words = []
        word_freq = {}
        
        for word, flag in words:
            word_lower = word.lower().strip()
            # 过滤条件：长度>1，不是纯停用词，词性符合要求
            if (len(word_lower) >= 2 and 
                word_lower not in self._stopwords and
                flag in allowed_pos):
                candidate_words.append(word_lower)
                word_freq[word_lower] = word_freq.get(word_lower, 0) + 1
        
        # 3. 如果没有有效的候选词，放宽条件：保留所有长度>=2且不是纯停用词的词
        if not candidate_words:
            words = jieba.cut(query)
            for word in words:
                word_lower = word.lower().strip()
                if len(word_lower) >= 2 and word_lower not in self._stopwords:
                    candidate_words.append(word_lower)
                    word_freq[word_lower] = word_freq.get(word_lower, 0) + 1
        
        # 4. 按词频排序，去重后返回
        unique_words = list(dict.fromkeys(candidate_words))  # 保持顺序去重
        # 优先返回高频词
        unique_words.sort(key=lambda x: word_freq.get(x, 0), reverse=True)
        
        keywords = unique_words[:MAX_KEYWORDS]
        print(f"🔑 规则提取的关键词: {keywords}")
        return keywords

    def _extract_keywords(self, query: str) -> list[str]:
        """
        统一的关键词提取接口
        根据配置选择使用LLM或规则提取
        """
        if USE_LLM_FOR_KEYWORDS:
            return self._extract_keywords_llm(query)
        else:
            return self._extract_keywords_rule_based(query)

    def _augment_query(self, query: str) -> tuple[str, list[str]]:
        """
        返回原始查询和提取的关键词列表
        """
        if not ENABLE_KEYWORD_AUGMENT:
            return query, []
        
        keywords = self._extract_keywords(query)
        print(f"📝 问题: {query}")
        print(f"🔑 提取关键词: {keywords}")
        
        # 可选：构建增强查询（将关键词拼接回去）
        # 但我们的多路召回会分别使用原始查询和关键词，所以这里只需要返回关键词
        return keywords
#===============query预处理=====================
    def _preprocess_query(self, query: str) -> str:
        """
        [OPT-02] 优化的Query预处理。

        原版问题：
        - 删除所有中文标点并把词拼在一起，破坏语义结构
        - 全角→半角转换方向错误
        - 停用词过滤后去掉所有空格，导致英文短语丢失边界

        优化策略：
        1. 保留核心标点和空格作为词语边界
        2. 仅过滤真正的停用词（保留技术术语）
        3. 保持中英文混排的自然结构
        4. 保留数字、单位、缩写
        """
        if not isinstance(query, str) or not query.strip():
            return ""

        original_query = query
        query = query.strip()

        # [OPT-02] 全角→半角转换（修复方向）
        query = self._fix_full_half_width(query)

        # 去除换行/制表符，保留空格
        query = query.replace("\n", " ").replace("\r", " ").replace("\t", " ")

        # [OPT-02] 仅去除极特殊字符，保留中英文标点和数字
        query = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9，。！？；：""''（）【】.,!?[]()-+=/%#&*]', ' ', query)

        # 去除常见问候语前缀
        greetings = [
            '你好', '您好', '请问', '我想问', '能不能', '可以告诉我', '麻烦问一下',
            '帮忙看看', '请教一下', '咨询一下', '了解一下', '想知道', '请问一下',
        ]
        for greeting in greetings:
            if query.lower().startswith(greeting.lower()):
                query = query[len(greeting):]
                break

        # 合并多余空格
        query = re.sub(r'\s{2,}', ' ', query).strip()

        # 如果清洗后为空或过短，返回原始查询
        if len(query.replace(' ', '')) < 2:
            return original_query.strip()

        return query

    def _fix_full_half_width(self, text: str) -> str:
        """
        [OPT-03] 修复：正确的全角→半角转换。

        原版bug：将半角空格(32)转全角空格(12288)，方向完全反了。
        修复：全角字符(FF01-FF5E)转半角(21-7E)，全角空格(3000)转半角空格(32)。
        """
        result = []
        for char in text:
            code = ord(char)
            if 0xFF01 <= code <= 0xFF5E:
                code -= 0xFEE0
            elif code == 0x3000:  # 全角空格
                code = 0x0020
            result.append(chr(code))
        return "".join(result)

    #================= 三路召回 ==================
    #================= 关键词匹配召回 ==================
    def _keyword_recall_with_augment(self, query: str, augmented_keywords: list = None) -> list:
        """
        增强版关键词召回：使用原始查询 + 增强关键词
        """
        # 如果有增强关键词，合并到查询中
        if augmented_keywords and ENABLE_KEYWORD_AUGMENT:
            # 将关键词拼接到查询中，增加权重
            enhanced_query = query + " " + " ".join(augmented_keywords)
            print(f"🔍 关键词召回增强查询: {enhanced_query}")
        else:
            enhanced_query = query
        
        tokenized_query = list(jieba.cut(enhanced_query.strip()))
        query_words = set(tokenized_query)
        
        # 过滤停用词
        query_words = {w for w in query_words if w not in self._stopwords}
        
        # 使用倒排索引快速计算
        doc_scores = {}
        for word in query_words:
            if word in self._keyword_index:
                for doc_idx in self._keyword_index[word]:
                    doc_scores[doc_idx] = doc_scores.get(doc_idx, 0) + 1
        
        scored = [(idx, score) for idx, score in doc_scores.items()]       
        scored.sort(key=lambda x: x[1], reverse=True)
        return scored[:TOP_K]

    def _keyword_recall(self, query: str, topk=TOP_K):
        tokenized_query = list(jieba.cut(query.strip())) #对查询语句做分词，拆成关键词列表
        query_words = set(tokenized_query) # 对关键词列表去重
        query_words = {w for w in query_words if w not in self._stopwords and len(w) > 1}
        
        if not query_words:
            return []
        # 使用倒排索引快速计算
        doc_scores = {}
        for word in query_words:
            if word in self._keyword_index:
                for doc_idx in self._keyword_index[word]:
                    doc_scores[doc_idx] = doc_scores.get(doc_idx, 0) + 1
        
        scored = [(idx, score) for idx, score in doc_scores.items()]
        scored.sort(key=lambda x: x[1], reverse=True) # 按匹配的关键词数量排序，从大到小
        return scored[:topk]

    #================= BM25召回 ==================
    def _bm25_recall_with_augment(self, query: str, augmented_keywords: list = None) -> list:
        """
        增强版BM25召回：使用原始查询 + 增强关键词
        """
        if augmented_keywords and ENABLE_KEYWORD_AUGMENT:
            # 将关键词拼接到查询中，增强BM25匹配
            enhanced_query = query + " " + " ".join(augmented_keywords)
            print(f"🔍 BM25召回增强查询: {enhanced_query}")
        else:
            enhanced_query = query
        
        tokenized_query = list(jieba.cut(enhanced_query.strip()))
        tokenized_query = [t for t in tokenized_query if t not in self._stopwords]
        if not tokenized_query:
            return []
        
        scores = self._bm25_index.get_scores(tokenized_query)
        ranked = np.argsort(scores)[::-1][:TOP_K]
        return [(int(i), float(scores[i])) for i in ranked if scores[i] > 0]
        
    def _bm25_recall(self, query: str, topk=TOP_K):
        
        # 对查询做分词
        tokenized_query = list(jieba.cut(query.strip()))
        tokenized_query = [t for t in tokenized_query if t not in self._stopwords]
        if not tokenized_query:
            return []
        scores = self._bm25_index.get_scores(tokenized_query)
        if len(scores) <= topk:
            ranked = np.argsort(scores)[::-1][:topk]
        else:
            top_indices = np.argpartition(scores, -topk)[-topk:]
            ranked = top_indices[np.argsort(scores[top_indices])[::-1]]
        return [(int(i), float(scores[i])) for i in ranked]

    #================= 向量检索召回 ==================
    def _vector_recall(self, query: str, topk=TOP_K):
        # [OPT-14] 直接调用向量库检索，避免每次创建retriever对象
        try:
            docs_scores = self.vector_store.similarity_search_with_relevance_scores(
                query, k=topk
            )
        except Exception:
            ret = self.vector_store.as_retriever(
                search_type="similarity_score_threshold",
                search_kwargs={"k": topk, "score_threshold": SIMILARITY_THRESHOLD},
            )
            docs_scores = [(doc, doc.metadata.get("score", 0.0)) for doc in ret.invoke(query)]

        result = []
        for doc, score in docs_scores:
            if score < SIMILARITY_THRESHOLD:
                continue
            idx = doc.metadata.get("chunk_index")
            if idx is not None and 0 <= idx < len(self.chunks):
                result.append((int(idx), float(score)))
        return result
    #================== RRF融合 ==================
    def _rrf_fuse(self, lists, k=60):
        scores = {} # 初始化一个空字典，用于存储每个文档块的RRF分数，key是文档块索引，value是RRF分数
        #print(f"融合前候选：{len(lists)} 条\n")
        for lst in lists:
            #遍历列表里的每一条记录，enumerate(lst)给每个元素带上索引
            # rank：从0开始的排名，第一名是0，第二名是1
            # idx：列表里文档块索引
            # _：忽略，因为这里只需要索引
            for rank, (idx, _) in enumerate(lst):  
                s = 1.0 / (rank + k) # RRF核心公式，k是平滑参数，防止除0错误
                # 把当前item的RRF分数累加到总字典里，如果idx存在字典里，分数就累加，否则默认是0。
                # 这里已经实现了去重，因为idx是文档的唯一ID，scores字典的key是唯一的，idx是字典里的key值
                scores[idx] = scores.get(idx, 0) + s 
        return sorted(scores.items(), key=lambda x: x[1], reverse=True) # 按RRF分数从高到低排序

    #================== 重排 ==================
    def _cross_encoder_rerank(self, query: str, fused_indices: list):
        # [OPT-14] 限制送入重排模型的候选数量，显著降低CrossEncoder耗时
        limited = fused_indices[:RERANK_CANDIDATE_LIMIT]
        candidates = [self.chunks[idx].page_content for idx, _ in limited]
        docs_obj = [self.chunks[idx] for idx, _ in limited]
        pairs = [[query, c] for c in candidates]
        scores = self.rerank_model.predict(pairs, batch_size=RERANK_BATCH_SIZE).tolist()
        reranked = sorted(zip(docs_obj, scores), key=lambda x: x[1], reverse=True)
        return reranked[:RERANK_TOP_N]

    @staticmethod
    def _char_bigram_similarity(text_a: str, text_b: str) -> float:
        """基于字符二元组的Jaccard相似度，用于快速去重"""
        def bigrams(text: str) -> set:
            snippet = text[:300]
            return {snippet[i:i + 2] for i in range(len(snippet) - 1)}

        bg_a, bg_b = bigrams(text_a), bigrams(text_b)
        if not bg_a or not bg_b:
            return 0.0
        union = bg_a | bg_b
        return len(bg_a & bg_b) / len(union) if union else 0.0

    # ==================== [OPT-07] Chunk去重 ====================
    def _deduplicate_chunks(self, docs: List[LangchainDocument]) -> List[LangchainDocument]:
        """
        对重排后的chunk列表去重。
        [OPT-14] 默认使用字符相似度快速去重，避免每个chunk都调用embedding API。
        """
        if not ENABLE_DEDUP or len(docs) <= 1:
            return docs

        if DEDUP_METHOD == "text":
            deduped = []
            for doc in docs:
                content = doc.page_content[:500]
                if any(
                    self._char_bigram_similarity(content, kept.page_content[:500]) > TEXT_DEDUP_THRESHOLD
                    for kept in deduped
                ):
                    continue
                deduped.append(doc)
            return deduped

        deduped = []
        seen_embeddings = []
        for doc in docs:
            is_dup = False
            try:
                doc_vec = np.array(self.embeddings.embed_query(doc.page_content[:500]))
            except Exception:
                doc_vec = None

            if doc_vec is not None:
                for seen_vec in seen_embeddings:
                    sim = float(np.dot(doc_vec, seen_vec) /
                                (np.linalg.norm(doc_vec) * np.linalg.norm(seen_vec) + 1e-8))
                    if sim > DEDUP_THRESHOLD:
                        is_dup = True
                        break
                if not is_dup:
                    seen_embeddings.append(doc_vec)

            if not is_dup:
                deduped.append(doc)

        return deduped

# ====================== 单Query检索 ======================
    def _single_query_retrieve(self, query: str) -> list:
        """单Query的三路召回"""
        augmented_keywords = self._augment_query(query)

        # [OPT-14] 三路召回并行执行（BM25/向量/关键词互不依赖）
        if ENABLE_PARALLEL_RECALL:
            with ThreadPoolExecutor(max_workers=3) as executor:
                f_bm25 = executor.submit(self._bm25_recall_with_augment, query, augmented_keywords)
                f_vec = executor.submit(self._vector_recall, query)
                f_kw = executor.submit(self._keyword_recall_with_augment, query, augmented_keywords)
                return [f_bm25.result(), f_vec.result(), f_kw.result()]

        bm25 = self._bm25_recall_with_augment(query, augmented_keywords)
        vec = self._vector_recall(query)
        kw = self._keyword_recall_with_augment(query, augmented_keywords)
        return [bm25, vec, kw]
# ====================== 多Query检索======================
    #=====================生成多个query=====================
    def _generate_multi_queries(self, original_query: str) -> list[str]:
        """
        基于原始查询生成多个不同角度的查询语句（多Query扩展）
        :param original_query: 用户原始问题
        :return: 包含原始查询的多个查询列表
        """
        if not ENABLE_MULTI_QUERY:
            return [original_query]
        
        # 多Query生成提示词模板
        multi_query_prompt = ChatPromptTemplate.from_messages([
            ("system", """你是查询扩展助手。根据用户问题生成{num}个不同角度的查询语句。
规则：
1. 保持核心语义，从不同表述/角度/关键词组合生成
2. 每个查询简洁，20字以内
3. 只输出查询语句，每行一个
4. 必须生成{num}个查询"""),
            ("human", "原始问题：{query}")
        ])
        try:
            chain = multi_query_prompt | self.multi_query_model
            response = chain.invoke({"query": original_query, "num": MULTI_QUERY_NUM})
            content = response.content if hasattr(response, 'content') else str(response)
            queries = [q.strip() for q in content.split('\n') if q.strip()]
            all_queries = list({original_query} | set(queries))[:MULTI_QUERY_NUM + 1]
            return all_queries
        except Exception:
            return [original_query]
    def _multi_query_retrieve(self, clean_query: str) -> list[LangchainDocument]:
        """
        多Query检索主流程：
        1. 生成多个扩展Query
        2. 每个Query独立执行三路召回
        3. 融合所有召回结果
        4. 重排得到最终结果
        """
        # 步骤1：生成多Query
        all_queries = self._generate_multi_queries(clean_query)
        
        # 步骤2：每个Query独立检索（多Query之间也并行）
        all_retrieval_results = []
        with ThreadPoolExecutor(max_workers=min(len(all_queries), MAX_WORKERS)) as executor:
            futures = [executor.submit(self._single_query_retrieve, q) for q in all_queries]
            for future in as_completed(futures):
                try:
                    all_retrieval_results.extend(future.result())
                except Exception:
                    pass
        
        # 步骤3：融合所有检索结果
        print(f"\n==== 融合所有Query的检索结果 ====")
        fused = self._rrf_fuse(all_retrieval_results)
        print(f"多Query融合后候选：{len(fused)} 条")

        # 步骤4：CrossEncoder重排（或 FAST_MODE 下直接取 RRF Top-N）
        reranked_docs = self._select_top_docs(clean_query, fused)
        return reranked_docs
    def _select_top_docs(self, query: str, fused: list) -> List[LangchainDocument]:
        """从重排或 RRF 结果中选取 Top-N 文档"""
        if not fused:
            return []
        if FAST_MODE or not ENABLE_RERANK:
            return [self.chunks[idx] for idx, _ in fused[:RERANK_TOP_N]]
        reranked = self._cross_encoder_rerank(query, fused)
        return [doc for doc, _ in reranked]

    def _retrieve_with_rerank(self, query: str) -> list[LangchainDocument]:
        """
        核心检索：预处理 → 三路召回 → RRF融合 → 重排/Top-N → 去重 → 截断
        """
        clean_query = self._preprocess_query(query)
        if not clean_query:
            return []

        if ENABLE_MULTI_QUERY:
            docs = self._multi_query_retrieve(clean_query)
        else:
            t0 = time.time()
            single_result = self._single_query_retrieve(clean_query)
            fused = self._rrf_fuse(single_result)
            if not fused:
                return []
            docs = self._select_top_docs(clean_query, fused)
            if ENABLE_PERF_LOG:
                print(f"  检索+融合+重排: {time.time() - t0:.2f}s")

        t1 = time.time()
        docs = self._deduplicate_chunks(docs)
        docs = self._truncate_by_token_budget(docs)
        if ENABLE_PERF_LOG:
            print(f"  去重+截断: {time.time() - t1:.2f}s")

        if ENABLE_DOC_COMPRESSION:
            docs = self._smart_compress_documents(docs, clean_query)

        return docs

    # ==================== [OPT-12] Token预算管理 ====================
    def _truncate_by_token_budget(self, docs: List[LangchainDocument]) -> List[LangchainDocument]:
        """
        按token预算截断检索结果，避免超出LLM上下文窗口。
        原理：LLM有最大context长度限制(默认4096 tokens≈3000中文)，
              超出部分会被截断，可能导致关键信息丢失。
              按相关性排序后优先保留高分chunk。
        """
        if not docs:
            return docs
        total_chars = sum(len(d.page_content) for d in docs)
        # 中文约1字符≈1.3token
        estimated_tokens = int(total_chars * 1.3)
        if estimated_tokens <= MAX_CONTEXT_TOKENS:
            return docs

        # 按顺序截断(已按相关性排序)
        budget_chars = int(MAX_CONTEXT_TOKENS / 1.3)
        kept = []
        used = 0
        for doc in docs:
            doc_len = len(doc.page_content)
            if used + doc_len <= budget_chars:
                kept.append(doc)
                used += doc_len
            else:
                remaining = budget_chars - used
                if remaining > 100:  # 至少保留100字符才有意义
                    truncated_doc = LangchainDocument(
                        page_content=doc.page_content[:remaining],
                        metadata={**doc.metadata, "truncated": True}
                    )
                    kept.append(truncated_doc)
                break
        return kept
    def _get_rag_prompt(self):
        """复用Prompt模板，避免重复构建"""
        if self._rag_prompt is None:
            self._rag_prompt = ChatPromptTemplate.from_messages([
                ("system", """你是一个专业的技术文档问答助手。请严格按以下规则回答：

【核心规则 - 必须严格遵守】
1. 只能基于【参考资料】中的内容回答，绝对禁止使用你自己的知识
2. 如果【参考资料】中没有相关信息，直接回答："根据所提供的资料，无法回答此问题。"
3. 禁止猜测、编造、推理、补充任何参考资料中没有的信息
4. 如果参考资料中有答案，请完整、准确地复述相关内容

【回答格式要求】
- 对于事实查询：简洁直接地给出答案，关键数据必须准确
- 对于总结归纳：分点列出要点，保持逻辑顺序
- 对于对比分析：先分别描述各方特点，再总结差异
- 如需引用资料中的具体数值或公式，请原样保留

【参考资料】
{context}"""),
                ("human", "{input}")
            ])
        return self._rag_prompt

    def _ensure_document_chain(self):
        """懒加载文档组合链，供生成阶段复用"""
        if self._document_chain is None:
            self._document_chain = create_stuff_documents_chain(
                self.chat_model, self._get_rag_prompt()
            )

    def _generate_answer(self, question: str, docs: List[LangchainDocument]) -> str:
        """基于已检索文档直接生成答案，跳过重复检索"""
        self._ensure_document_chain()
        response = self._document_chain.invoke({"input": question, "context": docs})
        return response if isinstance(response, str) else str(response)

    def _build_rag_chain(self):
        """构建 RAG 链（仅组装 Runnable，不加载 LLM，耗时 <1ms）"""
        if self.rag_chain is not None:
            return

        def retrieve(inputs):
            q = inputs["input"]
            docs = self._retrieve_with_rerank(q)
            return {"input": q, "context": docs}

        def run_chain(inputs):
            self._ensure_document_chain()
            return self._document_chain.invoke(inputs)

        self.rag_chain = RunnableLambda(retrieve) | RunnableLambda(run_chain)

    def ask_question(self, question: str) -> str:
        """向知识库提问"""
        if self.vector_store is None:
            return "错误：向量库未初始化，请先放入文档并构建向量库。"

        if not question or not question.strip():
            return "请输入有效的问题。"

        if self.rag_chain is None:
            try:
                self._build_rag_chain()
            except Exception as e:
                return f"错误：RAG链初始化失败 - {str(e)}"

        try:
            t0 = time.time()
            t_ret, t_gen = t0, t0
            if ENABLE_PERF_LOG:
                clean = self._preprocess_query(question)
                t_ret = time.time()
                docs = self._retrieve_with_rerank(clean or question)
                print(f"  检索: {time.time() - t_ret:.2f}s")
                t_gen = time.time()
                response = self._generate_answer(question, docs)
                print(f"  生成: {time.time() - t_gen:.2f}s")
                print(f"  总计: {time.time() - t0:.2f}s")
            else:
                response = self.rag_chain.invoke({"input": question})
            return response if isinstance(response, str) else str(response)
        except Exception as e:
            return f"错误：处理问题时发生异常 - {str(e)}"

    
    # ===================== 企业级文档管理 =====================
    # 核心原则：
    #   1. 基于文件哈希(MD5) + 修改时间 双重变更检测
    #   2. 使用确定性ID (source_path#chunk_N) 支持精确增删改
    #   3. 清单持久化，每次变更原子写入
    #   4. 操作加锁，防止并发冲突
    #   5. 变更后自动重建BM25和关键词索引

    @staticmethod
    def _compute_file_hash(file_path: str) -> str:
        """计算文件MD5哈希（大文件分块读取）"""
        md5 = hashlib.md5()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(8192), b''):
                md5.update(chunk)
        return md5.hexdigest()

    def _reload_chunks_from_vectorstore(self) -> None:
        """从向量库全量重载chunks列表（增量操作后保持一致性）"""
        all_docs_data = self.vector_store.get()
        self.chunks = []
        if all_docs_data and 'documents' in all_docs_data:
            for doc_id, text, metadata in zip(
                all_docs_data.get('ids', []),
                all_docs_data.get('documents', []),
                all_docs_data.get('metadatas', []),
            ):
                doc = LangchainDocument(
                    page_content=text,
                    metadata=metadata or {}
                )
                self.chunks.append(doc)

        for i, chunk in enumerate(self.chunks):
            chunk.metadata["chunk_index"] = i

    def _rebuild_indices(self) -> None:
        """重建BM25和关键词倒排索引（文档变更后调用）"""
        if self.chunks:
            self._build_bm25_index()
            self._build_keyword_index()
        else:
            self._bm25_index = None
            self._keyword_index = {}

    def _rebuild_manifest_from_chunks(self) -> None:
        """从现有chunks重建清单（用于首次建库或清单丢失时）"""
        source_groups: dict[str, list] = {}
        for i, chunk in enumerate(self.chunks):
            source = chunk.metadata.get('source', '')
            if source:
                source_groups.setdefault(source, []).append(i)

        for file_path, chunk_indices in source_groups.items():
            file_hash = ''
            last_modified = 0.0
            if os.path.exists(file_path):
                try:
                    file_hash = self._compute_file_hash(file_path)
                    last_modified = os.path.getmtime(file_path)
                except OSError:
                    pass

            chunk_ids = [f"{file_path}#chunk_{idx}" for idx in chunk_indices]
            self._manifest.add_record(DocumentRecord(
                file_path=file_path,
                file_hash=file_hash,
                last_modified=last_modified,
                chunk_ids=chunk_ids,
                chunk_count=len(chunk_indices),
                file_type=os.path.splitext(file_path)[1].lower(),
                indexed_at=datetime.now().isoformat(),
            ))

        self._manifest.save()
        print(f"清单已重建: {self._manifest.get_total_documents()} 个文件, "
              f"{self._manifest.get_total_chunks()} 个chunk")

    def add_document(self, file_path: str) -> bool:
        """
        企业级文档添加：
        1. 检测文件有效性
        2. 加载文档（支持txt/pdf/docx/doc）
        3. 切分为chunks，生成确定性ID
        4. 写入向量库
        5. 更新chunks列表和清单
        6. 重建检索索引
        """
        if not os.path.exists(file_path):
            print(f"文件不存在: {file_path}")
            return False

        abs_path = os.path.abspath(file_path)
        with self._mutation_lock:
            return self._add_document_locked(abs_path)

    def _add_document_locked(self, file_path: str) -> bool:
        """持有锁的文档添加实现"""
        file_hash = self._compute_file_hash(file_path)
        ext = os.path.splitext(file_path)[1].lower()

        if self._manifest.has_file(file_path):
            existing = self._manifest.get_record(file_path)
            if existing and existing.file_hash == file_hash:
                print(f"文档已存在且未变更，跳过: {os.path.basename(file_path)}")
                return True
            else:
                print(f"检测到文档已变更，先移除旧版本: {os.path.basename(file_path)}")
                self._remove_document_locked(file_path)

        try:
            if ext == '.txt':
                loader = TextLoader(file_path, encoding='utf-8')
                raw_docs = loader.load()
            elif ext == '.pdf':
                raw_docs = self._load_pdf_with_fallback(file_path)
            elif ext in ('.docx', '.doc'):
                doc = self._load_docx_with_heading(file_path)
                raw_docs = [doc]
            else:
                print(f"不支持的文件类型: {ext}")
                return False
        except Exception as e:
            print(f"加载文件失败 {file_path}: {e}")
            return False

        if not raw_docs:
            print(f"文件无有效内容: {file_path}")
            return False

        chunks = self._split_documents(raw_docs)
        if not chunks:
            print(f"切分后无有效文本块: {file_path}")
            return False

        current_max_idx = len(self.chunks)
        chunk_ids = []
        for i, chunk in enumerate(chunks):
            chunk_id = f"{file_path}#chunk_{current_max_idx + i}"
            chunk_ids.append(chunk_id)
            chunk.metadata["source"] = file_path
            chunk.metadata["chunk_index"] = current_max_idx + i

        try:
            self.vector_store.add_documents(chunks, ids=chunk_ids)
        except Exception as e:
            print(f"写入向量库失败: {e}")
            return False

        self._reload_chunks_from_vectorstore()

        self._manifest.add_record(DocumentRecord(
            file_path=file_path,
            file_hash=file_hash,
            last_modified=os.path.getmtime(file_path),
            chunk_ids=chunk_ids,
            chunk_count=len(chunks),
            file_type=ext,
            indexed_at=datetime.now().isoformat(),
        ))
        self._manifest.save()

        self._rebuild_indices()

        print(f"文档已添加: {os.path.basename(file_path)} ({len(chunks)} chunks)")
        return True

    def remove_document(self, file_path: str) -> bool:
        """
        企业级文档删除：
        1. 从向量库按source过滤删除所有相关chunks
        2. 从清单中移除记录
        3. 重载chunks列表
        4. 重建检索索引
        """
        abs_path = os.path.abspath(file_path)
        with self._mutation_lock:
            return self._remove_document_locked(abs_path)

    def _remove_document_locked(self, file_path: str) -> bool:
        """持有锁的文档删除实现"""
        if not self._manifest.has_file(file_path):
            print(f"文档不在知识库中: {os.path.basename(file_path)}")
            return False

        record = self._manifest.get_record(file_path)
        chunk_count = record.chunk_count if record else 0

        try:
            self.vector_store.delete(where={"source": file_path})
        except Exception as e:
            print(f"从向量库删除失败 {file_path}: {e}")
            return False

        self._manifest.remove_record(file_path)
        self._manifest.save()

        self._reload_chunks_from_vectorstore()
        self._rebuild_indices()

        print(f"文档已移除: {os.path.basename(file_path)} ({chunk_count} chunks)")
        return True

    def update_document(self, file_path: str) -> bool:
        """
        企业级文档更新（哈希变更检测 + 先删后加）：
        1. 检查文件是否存在及哈希是否变化
        2. 若已入库且未变化，跳过
        3. 若已入库但已变化，先删除旧版再添加新版
        4. 若未入库，直接添加
        """
        if not os.path.exists(file_path):
            print(f"文件不存在，尝试从知识库中移除: {file_path}")
            return self.remove_document(file_path)

        abs_path = os.path.abspath(file_path)
        with self._mutation_lock:
            if self._manifest.has_file(abs_path):
                current_hash = self._compute_file_hash(abs_path)
                existing = self._manifest.get_record(abs_path)
                if existing and existing.file_hash == current_hash:
                    print(f"文档未变更，跳过: {os.path.basename(abs_path)}")
                    return True
                self._remove_document_locked(abs_path)

            return self._add_document_locked(abs_path)

    def sync_knowledge_base(self) -> dict:
        """
        企业级知识库同步：
        扫描文档目录，与清单比对，批量增/删/改。

        返回统计字典: {added, removed, updated, skipped, errors}
        """
        stats = {'added': 0, 'removed': 0, 'updated': 0, 'skipped': 0, 'errors': 0}
        docs_path = Path(DOCS_DIR)

        if not docs_path.exists():
            print(f"文档目录不存在，创建: {DOCS_DIR}")
            docs_path.mkdir(parents=True)
            return stats

        with self._mutation_lock:
            current_files = set()
            for ext in ('*.txt', '*.pdf', '*.docx', '*.doc'):
                for f in docs_path.glob(ext):
                    current_files.add(str(f.resolve()))

            indexed_files = self._manifest.get_all_paths()

            removed_files = indexed_files - current_files
            for fp in sorted(removed_files):
                try:
                    if self._remove_document_locked(fp):
                        stats['removed'] += 1
                    else:
                        stats['errors'] += 1
                except Exception as e:
                    print(f"删除文档失败 {fp}: {e}")
                    stats['errors'] += 1

            for fp in sorted(current_files):
                try:
                    if self._manifest.has_file(fp):
                        current_hash = self._compute_file_hash(fp)
                        existing = self._manifest.get_record(fp)
                        if existing and existing.file_hash == current_hash:
                            stats['skipped'] += 1
                            continue
                        if self._remove_document_locked(fp):
                            if self._add_document_locked(fp):
                                stats['updated'] += 1
                            else:
                                stats['errors'] += 1
                        else:
                            stats['errors'] += 1
                    else:
                        if self._add_document_locked(fp):
                            stats['added'] += 1
                        else:
                            stats['errors'] += 1
                except Exception as e:
                    print(f"同步文档失败 {fp}: {e}")
                    stats['errors'] += 1

        print(f"\n知识库同步完成: "
              f"新增{stats['added']}, 移除{stats['removed']}, "
              f"更新{stats['updated']}, 跳过{stats['skipped']}, "
              f"失败{stats['errors']}")

        self._reload_chunks_from_vectorstore()
        self._rebuild_indices()

        return stats

    def get_document_status(self) -> dict:
        """获取知识库文档状态报告"""
        manifest_info = {
            'total_documents': self._manifest.get_total_documents(),
            'total_chunks': self._manifest.get_total_chunks(),
            'vector_count': self.vector_store._collection.count() if self.vector_store else 0,
        }

        documents = []
        for fp in sorted(self._manifest.get_all_paths()):
            rec = self._manifest.get_record(fp)
            exists = os.path.exists(fp)
            current_hash = ''
            if exists:
                try:
                    current_hash = self._compute_file_hash(fp)
                except OSError:
                    pass
            documents.append({
                'file': os.path.basename(fp),
                'path': fp,
                'type': rec.file_type if rec else '',
                'chunks': rec.chunk_count if rec else 0,
                'indexed_at': rec.indexed_at if rec else '',
                'on_disk': exists,
                'hash_match': rec.file_hash == current_hash if rec and exists else None,
            })

        return {
            'manifest': manifest_info,
            'documents': documents,
        }

    def start_file_watcher(self, poll_interval: float = 5.0) -> threading.Thread:
        """
        启动后台文件监控线程（基于轮询）。

        原理：定时扫描文档目录，检测文件增删改并自动同步到向量库。
        适用于生产环境中文档持续更新的场景。

        :param poll_interval: 轮询间隔（秒）
        :return: 监控线程对象，可调用 .join() 阻塞或 .stop() 停止
        """
        stop_event = threading.Event()

        def _watch_loop():
            print(f"文件监控已启动 (轮询间隔: {poll_interval}s, 目录: {DOCS_DIR})")
            last_state = {}
            docs_path = Path(DOCS_DIR)

            while not stop_event.is_set():
                try:
                    if not docs_path.exists():
                        stop_event.wait(poll_interval)
                        continue

                    current_state = {}
                    for ext in ('*.txt', '*.pdf', '*.docx', '*.doc'):
                        for f in docs_path.glob(ext):
                            fp = str(f.resolve())
                            try:
                                current_state[fp] = {
                                    'hash': self._compute_file_hash(fp),
                                    'mtime': os.path.getmtime(fp),
                                }
                            except OSError:
                                continue

                    if last_state:
                        current_set = set(current_state.keys())
                        last_set = set(last_state.keys())

                        added = current_set - last_set
                        removed = last_set - current_set
                        modified = {
                            fp for fp in (current_set & last_set)
                            if current_state[fp]['hash'] != last_state[fp].get('hash', '')
                        }

                        for fp in sorted(added):
                            print(f"[监控] 检测到新增文件: {os.path.basename(fp)}")
                            self.add_document(fp)

                        for fp in sorted(removed):
                            print(f"[监控] 检测到删除文件: {os.path.basename(fp)}")
                            self.remove_document(fp)

                        for fp in sorted(modified):
                            print(f"[监控] 检测到修改文件: {os.path.basename(fp)}")
                            self.update_document(fp)

                    last_state = current_state
                except Exception as e:
                    print(f"[监控] 扫描异常: {e}")

                stop_event.wait(poll_interval)

            print("文件监控已停止")

        watcher_thread = threading.Thread(target=_watch_loop, daemon=True, name="FileWatcher")
        watcher_thread._stop_event = stop_event
        watcher_thread.start()
        return watcher_thread

    def add_documents_batch(self, file_paths: str) -> dict:
        """
        批量添加文档（企业级批量处理）。

        优化：多个文档共享一次索引重建，而非每个文档重建一次。

        :param file_paths: 文档路径列表
        :return: 统计字典
        """
        stats = {'added': 0, 'skipped': 0, 'errors': 0}
        print(file_paths)
        with self._mutation_lock:
            for fp in os.listdir(os.path.join(file_paths)):
                print(fp)
                try:
                    result = self._add_document_locked(os.path.join(file_paths, fp))
                    if result:
                        stats['added'] += 1
                    else:
                        stats['errors'] += 1
                except Exception as e:
                    print(f"批量添加失败 {fp}: {e}")
                    stats['errors'] += 1

            self._reload_chunks_from_vectorstore()
            self._rebuild_indices()

        print(f"\n批量添加完成: 成功{stats['added']}, 跳过{stats['skipped']}, 失败{stats['errors']}")
        return stats

    def run_interactive(self):
        """交互式问答"""
        print("\n"+"="*55)
        print("欢迎使用本地知识库问答系统")
        print("="*55)
        print("输入问题与知识库对话，输入quit退出")
        while True:
            try:
                question = input("\n请输入问题：\n").strip()
                if not question:
                    continue
                if question.lower() == 'quit':
                    print("谢谢使用，再见！")
                    break

                print("检索中...", end="", flush=True)
                start_time = time.time()
                response = self.ask_question(question)
                elapsed = time.time() - start_time
                print(f"\n耗时: {elapsed:.2f}秒")
            except KeyboardInterrupt:
                print("\n\n已中断，输入 quit 退出。")
            except Exception as e:
                print(f"\n发生错误: {e}")

def main():
    parser = argparse.ArgumentParser(description="RAG知识库系统")
    parser.add_argument('--eval', action='store_true', help='运行评估模式')
    parser.add_argument('--testset', type=str, default=DEFAULT_TEST_SET, help='测试集路径')
    parser.add_argument('--limit', type=int, default=None, help='限制评估问题数量')
    parser.add_argument('--skip-build', action='store_true', help='跳过向量库构建')
    parser.add_argument('--no-warmup', action='store_true', help='跳过模型预热（首次问答会很慢）')
    parser.add_argument('--fast', action='store_true', help='快速模式：跳过重排，降低延迟')
    parser.add_argument('--perf', action='store_true', help='打印各阶段耗时明细')
    parser.add_argument('--ask', type=str, default=None, help='单次问答测试（例: --ask "什么是DCT"）')
    # 文档管理命令
    parser.add_argument('--add', type=str, default=None, help='添加单个文档到知识库')
    parser.add_argument('--remove', type=str, default=None, help='从知识库移除文档')
    parser.add_argument('--update', type=str, default=None, help='更新知识库中的文档')
    parser.add_argument('--sync', action='store_true', help='同步文档目录与知识库')
    parser.add_argument('--status', action='store_true', help='查看知识库文档状态')
    parser.add_argument('--watch', action='store_true', help='启动文件监控模式（自动同步文档变更）')
    parser.add_argument('--watch-interval', type=float, default=5.0, help='文件监控轮询间隔（秒）')
    parser.add_argument('--batch', type=str, default=None, help='批量添加文档目录（例: --batch ./docs）')

    args = parser.parse_args()
    if args.fast:
        global FAST_MODE, ENABLE_RERANK
        FAST_MODE = True
        ENABLE_RERANK = False
    if args.perf:
        global ENABLE_PERF_LOG
        ENABLE_PERF_LOG = True

    rag = LocalRAGSystem()
    if not args.skip_build:
        rag.build_or_load_db()

    # 文档管理命令（无需向量库即可执行status）
    if args.status:
        if rag.vector_store is None:
            print("向量库未初始化")
        else:
            status = rag.get_document_status()
            print("\n" + "=" * 55)
            print("  知识库文档状态")
            print("=" * 55)
            m = status['manifest']
            print(f"  已索引文件: {m['total_documents']}")
            print(f"  已索引chunk: {m['total_chunks']}")
            print(f"  向量库向量数: {m['vector_count']}")
            print("-" * 55)
            for doc in status['documents']:
                hash_icon = "✓" if doc['hash_match'] else ("✗" if doc['hash_match'] is False else "?")
                disk_icon = "✓" if doc['on_disk'] else "✗"
                print(f"  [{hash_icon}] disk={disk_icon} {doc['file']} ({doc['chunks']} chunks)")
            print("=" * 55)
        return

    if rag.vector_store is None and not args.add:
        print("向量库为空，请先放入文档并运行，或使用 --add 添加文档")
        return

    if args.add:
        rag.add_document(args.add)
        return

    if args.remove:
        rag.remove_document(args.remove)
        return

    if args.update:
        rag.update_document(args.update)
        return

    if args.sync:
        rag.sync_knowledge_base()
        return

    if args.watch:
        if not args.no_warmup:
            rag.warmup_models(preload_chat=True)
        print("=" * 55)
        print("  文件监控模式")
        print(f"  监控目录: {DOCS_DIR}")
        print("  按 Ctrl+C 停止监控")
        print("=" * 55)
        watcher = rag.start_file_watcher(poll_interval=args.watch_interval)
        try:
            while True:
                time.sleep(1)
        except KeyboardInterrupt:
            print("\n正在停止文件监控...")
            watcher._stop_event.set()
            watcher.join(timeout=10)
        return
    if args.batch:
        rag.add_documents_batch(args.batch)
        return
    if not args.no_warmup:
        rag.warmup_models(preload_chat=True)

    if args.ask:
        import time as _time
        t0 = _time.time()
        print(rag.ask_question(args.ask))
        print(f"\n单次问答总耗时: {_time.time() - t0:.1f}s")
        return

    if args.eval:
        print("n"+"="*55)
        print("进入评估模式"+"="*55)
        evaluator = RAGEvaluator(rag)
        summary = evaluator.run_evaluation(args.testset, limit=args.limit)
        print("\n评估完成！")
    else:
        rag.run_interactive()
    
if __name__ == '__main__':
    main()
